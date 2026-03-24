#!/usr/bin/env python3
import json
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "workspace" / "output"
LOGS_DIR = BASE_DIR / "workspace" / "logs"


ALIGNMENT_MAP = {
    "LEFT (0)": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER (1)": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT (2)": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY (3)": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    LOGS_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def write_log(message: str) -> None:
    log_path = LOGS_DIR / "reconstruct_docx.log"
    with log_path.open("a", encoding="utf-8") as f:
        f.write(message + "\n")


def safe_set_run_format(run, formatting: Dict[str, Any]) -> None:
    if not formatting:
        return

    if formatting.get("bold") is not None:
        run.bold = formatting.get("bold")

    if formatting.get("italic") is not None:
        run.italic = formatting.get("italic")

    if formatting.get("underline") is not None:
        run.underline = formatting.get("underline")

    if formatting.get("superscript") is not None:
        run.font.superscript = formatting.get("superscript")

    if formatting.get("subscript") is not None:
        run.font.subscript = formatting.get("subscript")

    if formatting.get("small_caps") is not None:
        run.font.small_caps = formatting.get("small_caps")

    if formatting.get("all_caps") is not None:
        run.font.all_caps = formatting.get("all_caps")

    if formatting.get("strike") is not None:
        run.font.strike = formatting.get("strike")

    if formatting.get("double_strike") is not None:
        run.font.double_strike = formatting.get("double_strike")

    if formatting.get("hidden") is not None:
        run.font.hidden = formatting.get("hidden")

    if formatting.get("name"):
        run.font.name = formatting.get("name")

    if formatting.get("size_pt") is not None:
        try:
            run.font.size = Pt(float(formatting.get("size_pt")))
        except Exception:
            pass


def apply_paragraph_style(paragraph, block: Dict[str, Any]) -> None:
    style_name = block.get("style_name")
    if style_name:
        try:
            paragraph.style = style_name
        except Exception:
            pass

    alignment_raw = (
        block.get("paragraph_format", {}) or {}
    ).get("alignment")

    if alignment_raw in ALIGNMENT_MAP:
        paragraph.alignment = ALIGNMENT_MAP[alignment_raw]


def choose_heading_level(block: Dict[str, Any]) -> Optional[int]:
    semantic_type = block.get("semantic_type")
    if semantic_type == "chapter_title":
        return 1
    if semantic_type == "section_title":
        return 2
    if semantic_type == "subsection_title":
        return 3

    level = block.get("heading_level")
    if isinstance(level, int):
        return level if level > 0 else 1

    return None


def distribute_text_across_runs(final_text: str, runs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Reparte el texto final en runs intentando conservar la estructura de estilos del original.
    Estrategia MVP:
    - si no hay runs, se crea uno solo
    - si hay un solo run, se asigna todo al run
    - si hay múltiples runs, se distribuye proporcionalmente por longitud original
    """
    if not runs:
        return [{
            "text": final_text,
            "formatting": {},
        }]

    original_lengths = [len((r.get("text") or "")) for r in runs]
    total_original = sum(original_lengths)

    if len(runs) == 1 or total_original == 0:
        new_runs = []
        first = deepcopy(runs[0])
        first["text"] = final_text
        new_runs.append(first)
        return new_runs

    new_runs: List[Dict[str, Any]] = []
    remaining_text = final_text
    assigned = 0

    for idx, run in enumerate(runs):
        run_copy = deepcopy(run)
        if idx == len(runs) - 1:
            run_copy["text"] = remaining_text
            new_runs.append(run_copy)
            break

        proportion = original_lengths[idx] / total_original if total_original else 0
        slice_len = round(len(final_text) * proportion)

        # Evitar vaciar demasiado el resto
        max_len = max(0, len(remaining_text) - (len(runs) - idx - 1))
        slice_len = min(slice_len, max_len)

        piece = remaining_text[:slice_len]
        remaining_text = remaining_text[slice_len:]
        run_copy["text"] = piece
        new_runs.append(run_copy)
        assigned += len(piece)

    return new_runs


def add_runs_to_paragraph(paragraph, block: Dict[str, Any], final_text: str) -> None:
    original_runs = block.get("runs", []) or []
    reconstructed_runs = distribute_text_across_runs(final_text, original_runs)

    if not reconstructed_runs:
        paragraph.add_run(final_text)
        return

    for run_data in reconstructed_runs:
        text = run_data.get("text", "")
        run = paragraph.add_run(text)
        safe_set_run_format(run, run_data.get("formatting", {}) or {})


def add_paragraph_block(document: Document, block: Dict[str, Any]) -> None:
    final_text = block.get("final_text", "")
    heading_level = choose_heading_level(block)

    if heading_level:
        paragraph = document.add_heading(level=min(max(heading_level, 1), 9))
        if paragraph.runs:
            # limpiamos el run inicial vacío si existe
            paragraph.runs[0].text = ""
    else:
        paragraph = document.add_paragraph()

    apply_paragraph_style(paragraph, block)
    add_runs_to_paragraph(paragraph, block, final_text)


def infer_table_dimensions(rows: List[List[Dict[str, Any]]]) -> tuple[int, int]:
    row_count = len(rows)
    col_count = max((len(row) for row in rows), default=0)
    return row_count, col_count


def add_table_block(document: Document, block: Dict[str, Any]) -> None:
    rows = block.get("rows", []) or []
    row_count, col_count = infer_table_dimensions(rows)

    if row_count == 0 or col_count == 0:
        document.add_paragraph("")
        return

    table = document.add_table(rows=row_count, cols=col_count)

    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            target_cell = table.cell(r_idx, c_idx)

            # Limpiar el párrafo por defecto
            if target_cell.paragraphs:
                target_cell.paragraphs[0].text = ""

            paragraphs = cell.get("paragraphs", []) or []
            if not paragraphs:
                target_cell.text = cell.get("final_text", "") or ""
                continue

            first_para = True
            for p in paragraphs:
                if first_para:
                    para = target_cell.paragraphs[0]
                    first_para = False
                else:
                    para = target_cell.add_paragraph()

                para_text = p.get("final_text", "") or ""
                para.style = p.get("style_name") or para.style
                runs = p.get("runs", []) or []
                reconstructed_runs = distribute_text_across_runs(para_text, runs)

                for run_data in reconstructed_runs:
                    run = para.add_run(run_data.get("text", ""))
                    safe_set_run_format(run, run_data.get("formatting", {}) or {})


def reconstruct_document(data: Dict[str, Any]) -> Document:
    document = Document()
    blocks = data.get("blocks", []) or []

    for block in blocks:
        if block.get("kind") == "table":
            add_table_block(document, block)
        else:
            add_paragraph_block(document, block)

    return document


def build_output_path(input_json_path: Path, data: Dict[str, Any]) -> Path:
    source_file = (data.get("document_metadata", {}) or {}).get("source_file")
    if source_file:
        stem = Path(source_file).stem
    else:
        stem = input_json_path.stem.replace(".validated", "")

    return OUTPUT_DIR / f"{stem}.translated.en.docx"


def main() -> int:
    ensure_dirs()

    if len(sys.argv) < 2:
        print("Usage: python reconstruct_docx.py <validated.json>")
        return 1

    input_arg = Path(sys.argv[1])
    input_path = input_arg.resolve() if input_arg.is_absolute() else (Path.cwd() / input_arg).resolve()

    if not input_path.exists():
        print(f"Error: file not found: {input_path}")
        return 1

    if input_path.suffix.lower() != ".json":
        print("Error: input file must be a .json")
        return 1

    try:
        data = load_json(input_path)
        document = reconstruct_document(data)
        output_path = build_output_path(input_path, data)
        document.save(str(output_path))

        validation_summary = data.get("validation_summary", {}) or {}
        write_log(
            f"OK | reconstructed={input_path} | output={output_path} "
            f"| validation_status={validation_summary.get('status')}"
        )

        print(f"DOCX reconstruction complete: {output_path}")
        return 0
    except Exception as exc:
        write_log(f"ERROR | file={input_path} | error={repr(exc)}")
        print(f"DOCX reconstruction failed: {exc}")
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
